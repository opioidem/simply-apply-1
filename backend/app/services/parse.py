"""Resume upload → text → structured JSON Resume.

Two stages, deliberately separated:

  1. Extraction (deterministic): PDF/DOCX/TXT → plain text. No model involved, so it
     either works or fails loudly.
  2. Structuring (LLM): text → StructuredResume, schema-constrained by the provider.

The parse is never trusted. The API returns it for the user to review and correct on a
confirm screen before anything is stored as the base resume — every later render and
every guardrail check is anchored to that confirmed document, so an unreviewed parse
error would silently propagate into every application.
"""

from __future__ import annotations

import io
import logging

from app.llm.base import LLMProvider
from app.schemas import StructuredResume

log = logging.getLogger(__name__)

MAX_RESUME_CHARS = 30000

SYSTEM_PROMPT = """Extract the contents of a resume into structured JSON Resume format.

Transcribe only. Do not summarize, improve, correct, or infer:
- Copy dates exactly as written ("2021 - Present", "Jun 2019").
- Copy every bullet point verbatim, including numbers and percentages.
- If a field is absent, leave it empty. Never guess at an email, a date, or a title.
- Preserve the original ordering of entries.

This becomes the single source of truth the user's future resumes are generated from, \
so a transcription error here propagates into every application they send. Accuracy \
matters more than tidiness."""


class ParseError(RuntimeError):
    """Raised when a file cannot be read as text."""


def extract_text(filename: str, data: bytes) -> str:
    """Deterministic file → text. Raises ParseError on unsupported or unreadable input."""
    name = (filename or "").lower()

    if name.endswith(".pdf"):
        text = _extract_pdf(data)
    elif name.endswith(".docx"):
        text = _extract_docx(data)
    elif name.endswith((".txt", ".md")):
        text = data.decode("utf-8", errors="replace")
    elif name.endswith(".doc"):
        raise ParseError(
            "Legacy .doc files aren't supported. Save as .docx or PDF and try again."
        )
    else:
        raise ParseError(
            f"Unsupported file type: {filename!r}. Upload a PDF, DOCX, TXT, or MD file."
        )

    text = text.strip()
    if not text:
        raise ParseError(
            "No text could be extracted. If this is a scanned or image-only PDF, it needs "
            "OCR first — SimplyApply doesn't do image extraction."
        )
    return text[:MAX_RESUME_CHARS]


def _extract_pdf(data: bytes) -> str:
    try:
        import pdfplumber
    except ImportError as exc:  # pragma: no cover - dependency guard
        raise ParseError("pdfplumber is not installed; cannot read PDF files.") from exc

    try:
        pages: list[str] = []
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages:
                pages.append(page.extract_text() or "")
        return "\n\n".join(pages)
    except Exception as exc:
        raise ParseError(f"Could not read the PDF: {exc}") from exc


def _extract_docx(data: bytes) -> str:
    try:
        import docx
    except ImportError as exc:  # pragma: no cover - dependency guard
        raise ParseError("python-docx is not installed; cannot read DOCX files.") from exc

    try:
        document = docx.Document(io.BytesIO(data))
        blocks = [p.text for p in document.paragraphs]
        # Many resume templates lay content out in tables; paragraphs alone would miss it.
        for table in document.tables:
            for row in table.rows:
                blocks.extend(cell.text for cell in row.cells)
        return "\n".join(b for b in blocks if b and b.strip())
    except Exception as exc:
        raise ParseError(f"Could not read the DOCX: {exc}") from exc


async def parse_resume(provider: LLMProvider, text: str) -> StructuredResume:
    """Structure extracted resume text. The caller must have the user confirm the result."""
    return await provider.complete_structured(
        system=SYSTEM_PROMPT,
        user=f"RESUME TEXT\n{text}",
        schema=StructuredResume,
    )
