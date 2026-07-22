"""DOCX renderer — the primary ATS output.

ATS parsers are the audience, not humans. That constrains the layout more than it might
seem:

  * Single column, no tables, no text boxes. Multi-column layouts are the single most
    common cause of a resume being read in the wrong order — a parser walks the XML in
    document order, so a two-column layout interleaves your job titles with your skills.
  * No headers or footers. Many parsers skip them entirely, so contact details placed
    there vanish.
  * Standard section headings ("Experience", "Education", "Skills"). Parsers match on
    these literal words; "What I've Been Up To" is invisible to them.
  * A real bullet style, not a "•" character typed into the paragraph text.

Pure python-docx, no system dependencies — this path works everywhere, which is why the
PRD makes it the primary output rather than the PDF.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from app.schemas import StructuredResume

BODY_FONT = "Calibri"  # metrically standard, present on every platform, parses cleanly
BODY_SIZE = Pt(10)
NAME_SIZE = Pt(20)
HEADING_SIZE = Pt(12)
INK = RGBColor(0x26, 0x31, 0x37)


def _configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = BODY_SIZE
    normal.font.color.rgb = INK
    paragraph_format = normal.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(2)


def _section_heading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text.upper())
    run.bold = True
    run.font.size = HEADING_SIZE
    run.font.color.rgb = INK


def _date_range(start: str, end: str) -> str:
    if start and end:
        return f"{start} – {end}"
    return start or end or ""


def _contact_line(resume: StructuredResume) -> str:
    basics = resume.basics
    location = ", ".join(
        p for p in (basics.location.city, basics.location.region) if p
    )
    pieces = [basics.email, basics.phone, location, basics.url]
    pieces += [p.url for p in basics.profiles if p.url]
    # Plain separator: some parsers treat "|" as a column delimiter and split the line.
    return "  ·  ".join(p for p in pieces if p)


def render_docx(resume: StructuredResume, output_path: Path) -> Path:
    document = Document()
    _configure_styles(document)

    for section in document.sections:
        section.top_margin = Pt(40)
        section.bottom_margin = Pt(40)
        section.left_margin = Pt(50)
        section.right_margin = Pt(50)

    # --- header (in the body, never in a real header) --------------------
    name_paragraph = document.add_paragraph()
    name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_paragraph.paragraph_format.space_after = Pt(2)
    name_run = name_paragraph.add_run(resume.basics.name or "")
    name_run.bold = True
    name_run.font.size = NAME_SIZE
    name_run.font.color.rgb = INK

    if resume.basics.label:
        label_paragraph = document.add_paragraph()
        label_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label_run = label_paragraph.add_run(resume.basics.label)
        label_run.font.size = Pt(11)

    contact = _contact_line(resume)
    if contact:
        contact_paragraph = document.add_paragraph()
        contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_paragraph.paragraph_format.space_after = Pt(6)
        contact_paragraph.add_run(contact).font.size = Pt(9)

    # --- summary ---------------------------------------------------------
    if resume.basics.summary:
        _section_heading(document, "Summary")
        document.add_paragraph(resume.basics.summary)

    # --- experience ------------------------------------------------------
    if resume.work:
        _section_heading(document, "Experience")
        for job in resume.work:
            line = document.add_paragraph()
            line.paragraph_format.space_before = Pt(6)
            line.paragraph_format.space_after = Pt(0)
            if job.position:
                line.add_run(job.position).bold = True
            if job.name:
                line.add_run(f"{' — ' if job.position else ''}{job.name}")

            meta = "  ·  ".join(
                p for p in (_date_range(job.startDate, job.endDate), job.location) if p
            )
            if meta:
                meta_paragraph = document.add_paragraph()
                meta_paragraph.paragraph_format.space_after = Pt(2)
                meta_paragraph.add_run(meta).font.size = Pt(9)

            if job.summary:
                document.add_paragraph(job.summary)
            for highlight in job.highlights:
                document.add_paragraph(highlight, style="List Bullet")

    # --- education -------------------------------------------------------
    if resume.education:
        _section_heading(document, "Education")
        for edu in resume.education:
            line = document.add_paragraph()
            line.paragraph_format.space_before = Pt(4)
            line.paragraph_format.space_after = Pt(0)
            degree = " ".join(p for p in (edu.studyType, edu.area) if p)
            if degree:
                line.add_run(degree).bold = True
            if edu.institution:
                line.add_run(f"{' — ' if degree else ''}{edu.institution}")

            meta = "  ·  ".join(
                p
                for p in (
                    _date_range(edu.startDate, edu.endDate),
                    f"GPA {edu.score}" if edu.score else "",
                )
                if p
            )
            if meta:
                meta_paragraph = document.add_paragraph()
                meta_paragraph.add_run(meta).font.size = Pt(9)
            if edu.courses:
                document.add_paragraph(f"Relevant coursework: {', '.join(edu.courses)}")

    # --- skills ----------------------------------------------------------
    if resume.skills:
        _section_heading(document, "Skills")
        # A skill entry is either a labelled group ("Languages: Python, SQL") or a bare
        # skill with nothing under it — resumes that list skills as flat bullets parse
        # into the latter. Emitting "name: " unconditionally leaves a dangling colon on
        # every bare entry, which looks broken in the finished document.
        for skill in resume.skills:
            paragraph = document.add_paragraph()
            if skill.name and skill.keywords:
                paragraph.add_run(f"{skill.name}: ").bold = True
                paragraph.add_run(", ".join(skill.keywords))
            elif skill.name:
                paragraph.add_run(skill.name)
            elif skill.keywords:
                paragraph.add_run(", ".join(skill.keywords))

    # --- projects --------------------------------------------------------
    if resume.projects:
        _section_heading(document, "Projects")
        for project in resume.projects:
            line = document.add_paragraph()
            line.paragraph_format.space_before = Pt(4)
            line.paragraph_format.space_after = Pt(0)
            if project.name:
                line.add_run(project.name).bold = True
            date_range = _date_range(project.startDate, project.endDate)
            if date_range:
                line.add_run(f"  ({date_range})").font.size = Pt(9)

            if project.description:
                document.add_paragraph(project.description)
            for highlight in project.highlights:
                document.add_paragraph(highlight, style="List Bullet")
            if project.keywords:
                document.add_paragraph(f"Technologies: {', '.join(project.keywords)}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    return output_path
