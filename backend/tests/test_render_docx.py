"""DOCX render tests — the actual ATS-survivability check.

The claim in the PRD is "extracts as clean, correctly-ordered text in a parser test."
These tests are that parser test: read the generated .docx back the way a naive ATS
would (document order, paragraph by paragraph) and assert sections come out in the right
order with content intact.
"""

from __future__ import annotations

import docx
import pytest

from app.schemas import Basics, Education, Project, Skill, StructuredResume, Work
from app.services.render_docx import render_docx


@pytest.fixture
def resume() -> StructuredResume:
    return StructuredResume(
        basics=Basics(
            name="Jane Doe",
            label="Backend Engineer",
            email="jane@example.com",
            phone="+1 555 0100",
            summary="Backend engineer focused on Python services.",
        ),
        work=[
            Work(
                name="Acme Corp",
                position="Software Engineer",
                startDate="2022-01",
                endDate="2024-06",
                highlights=["Reduced p95 latency by 15%.", "Migrated 12 services."],
            ),
            Work(name="Beta LLC", position="Intern", startDate="2021-06", endDate="2021-09"),
        ],
        education=[Education(institution="State University", studyType="BSc", area="Computer Science")],
        skills=[Skill(name="Languages", keywords=["Python", "SQL"])],
        projects=[Project(name="Ledger", description="Double-entry bookkeeping.")],
    )


def _extract(path) -> list[str]:
    """Read paragraphs in document order — what a parser walking the XML sees."""
    document = docx.Document(str(path))
    return [p.text.strip() for p in document.paragraphs if p.text.strip()]


def test_sections_appear_in_reading_order(resume, tmp_path) -> None:
    lines = _extract(render_docx(resume, tmp_path / "r.docx"))
    headings = [l for l in lines if l in {"SUMMARY", "EXPERIENCE", "EDUCATION", "SKILLS", "PROJECTS"}]
    assert headings == ["SUMMARY", "EXPERIENCE", "EDUCATION", "SKILLS", "PROJECTS"]


def test_name_is_first_line(resume, tmp_path) -> None:
    """Contact details in a real header get skipped by many parsers; ours is in the body."""
    assert _extract(render_docx(resume, tmp_path / "r.docx"))[0] == "Jane Doe"


def test_contact_details_survive(resume, tmp_path) -> None:
    joined = "\n".join(_extract(render_docx(resume, tmp_path / "r.docx")))
    assert "jane@example.com" in joined
    assert "+1 555 0100" in joined


def test_work_entries_keep_their_order(resume, tmp_path) -> None:
    lines = _extract(render_docx(resume, tmp_path / "r.docx"))
    joined = "\n".join(lines)
    assert joined.index("Acme Corp") < joined.index("Beta LLC")


def test_bullets_stay_with_their_employer(resume, tmp_path) -> None:
    """Interleaving is the classic multi-column failure — assert it doesn't happen."""
    joined = "\n".join(_extract(render_docx(resume, tmp_path / "r.docx")))
    acme = joined.index("Acme Corp")
    beta = joined.index("Beta LLC")
    assert acme < joined.index("Reduced p95 latency by 15%.") < beta


def test_metrics_are_not_mangled(resume, tmp_path) -> None:
    joined = "\n".join(_extract(render_docx(resume, tmp_path / "r.docx")))
    assert "15%" in joined
    assert "12 services" in joined


def test_no_tables_are_emitted(resume, tmp_path) -> None:
    """Tables are a top cause of out-of-order extraction. We must emit none."""
    document = docx.Document(str(render_docx(resume, tmp_path / "r.docx")))
    assert len(document.tables) == 0


def test_bare_skill_renders_without_dangling_colon(tmp_path) -> None:
    """Resumes that list skills as flat bullets parse to name-only entries.

    Emitting "name: " unconditionally left "Python: " with nothing after it on every
    such line — visible in the finished document, and only caught by reading a real
    generated file rather than asserting on the model.
    """
    resume = StructuredResume(
        skills=[Skill(name="Python"), Skill(name="Languages", keywords=["SQL", "Go"])]
    )
    lines = _extract(render_docx(resume, tmp_path / "s.docx"))
    assert "Python" in lines
    assert "Python:" not in lines
    assert "Languages: SQL, Go" in lines


def test_empty_resume_does_not_crash(tmp_path) -> None:
    path = render_docx(StructuredResume(), tmp_path / "empty.docx")
    assert path.exists()
